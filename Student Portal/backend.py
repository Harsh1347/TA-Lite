import PyPDF2
from docx import Document
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_ollama import ChatOllama
from langchain.schema import Document as LangchainDocument
import json
from langchain.chains import RetrievalQA

class LectureContent:
    def __init__(self, lecture_num: str):
        """Initialize with lecture number (01-10)"""
        self.lecture_num = lecture_num
        self.pdf_path = f"materials/LLMs-{lecture_num}.pdf"
        self.docx_path = f"transcript/transcript-{lecture_num}.docx"
        self.combined_text = None
        self.vector_store = None

    def load_content(self) -> bool:
        """Load and combine content from both PDF and DOCX files."""
        try:
            # Extract text from PDF (Summary)
            pdf_text = self._extract_text_from_pdf()
            print("\n=== Extracted Text from PDF ===")
            print(pdf_text[:500] + "..." if len(pdf_text) > 500 else pdf_text)
            print("=== End of PDF Text ===\n")

            # Extract text from DOCX (Notes)
            docx_text = self._extract_text_from_docx()
            print("\n=== Extracted Text from DOCX ===")
            print(docx_text[:500] + "..." if len(docx_text) > 500 else docx_text)
            print("=== End of DOCX Text ===\n")

            # Combine texts with clear separation
            self.combined_text = f"""=== Summary (PDF) ===\n{pdf_text}\n\n=== Detailed Notes (DOCX) ===\n{docx_text}"""
            
            # Prepare vector store
            self.vector_store = self._prepare_vector_store()
            
            return True
        except Exception as e:
            print(f"Error loading lecture content: {str(e)}")
            return False

    def _extract_text_from_pdf(self) -> str:
        """Extract text from PDF file."""
        try:
            with open(self.pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            print(f"Error extracting PDF text: {str(e)}")
            return ""

    def _extract_text_from_docx(self) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(self.docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error extracting DOCX text: {str(e)}")
            return ""

    def _prepare_vector_store(self):
        """Prepare vector store from combined text."""
        try:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=50,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
            chunks = text_splitter.split_text(self.combined_text)
            documents = [LangchainDocument(page_content=chunk) for chunk in chunks]
            embeddings = OllamaEmbeddings(model="llama2")
            return FAISS.from_documents(documents, embeddings)
        except Exception as e:
            print(f"Error preparing vector store: {str(e)}")
            return None

    def get_summary(self) -> str:
        """Generate a concise summary of the lecture content."""
        try:
            if not self.vector_store:
                return "Please load the lecture content first."
            
            llm = ChatOllama(model="llama2", temperature=0.7)
            
            # Get relevant content chunks
            all_docs = self.vector_store.similarity_search("", k=10)
            content = "\n".join([doc.page_content for doc in all_docs])
            
            prompt = f"""
            Generate a concise summary of Lecture {self.lecture_num}. The summary should:
            1. Highlight the main topics and key concepts covered
            2. Be well-structured and easy to understand
            3. Be approximately 3-4 paragraphs long
            4. Focus on the most important takeaways
            
            Content:
            {content}
            """
            
            response = llm.invoke(prompt)
            return response.content.strip()
            
        except Exception as e:
            print(f"Error generating summary: {str(e)}")
            return f"Error: {str(e)}"

    def get_study_notes(self) -> str:
        """Generate structured study notes from the lecture content."""
        try:
            if not self.vector_store:
                return "Please load the lecture content first."
            
            llm = ChatOllama(model="llama2", temperature=0.7)
            
            # Get relevant content chunks
            all_docs = self.vector_store.similarity_search("", k=10)
            content = "\n".join([doc.page_content for doc in all_docs])
            
            prompt = f"""
            Create clear and structured study notes for Lecture {self.lecture_num}. The notes should:
            1. Start with a brief overview of the lecture
            2. Be organized with clear headings and subheadings
            3. Use bullet points for key concepts and definitions
            4. Include relevant examples and explanations
            5. Highlight important terms and concepts
            6. End with key takeaways
            
            Content:
            {content}
            """
            
            response = llm.invoke(prompt)
            return response.content.strip()
            
        except Exception as e:
            print(f"Error generating study notes: {str(e)}")
            return f"Error: {str(e)}"

    def get_ai_response(self, question: str) -> str:
        """Get AI response for a question about the lecture content."""
        try:
            if not self.vector_store:
                return "Please load the lecture content first."
            
            llm = ChatOllama(model="llama2", temperature=0.7)
            
            qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=self.vector_store.as_retriever(search_kwargs={"k": 3}),
                return_source_documents=True
            )
            
            context = self.vector_store.similarity_search(question, k=3)
            context_text = "\n".join([doc.page_content for doc in context])
            
            print(f"\n=== Question about Lecture {self.lecture_num} ===")
            print(f"Question: {question}")
            print("\nContext used:")
            print(context_text[:500] + "..." if len(context_text) > 500 else context_text)
            print("=== End of Context ===\n")
            
            prompt = f"""
            You are a teaching assistant for Lecture {self.lecture_num}. Based on the following context and question, provide a clear and informative response.
            
            Context from lecture material:
            {context_text}
            
            Student Question:
            {question}
            
            Please provide a helpful response that:
            1. Directly addresses the question
            2. Uses information from the lecture content
            3. Explains concepts clearly
            4. Provides examples if relevant
            """
            
            response = qa_chain({"query": prompt})
            return response['result'].strip() if response and 'result' in response else "Sorry, I couldn't generate a response."
            
        except Exception as e:
            print(f"Error getting AI response: {str(e)}")
            return f"Error: {str(e)}"

def extract_text_from_file(file_path: str) -> str:
    """Extract text from PDF or DOCX file."""
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                print("\n=== Extracted Text from PDF ===")
                print(text[:500] + "..." if len(text) > 500 else text)
                print("=== End of Extracted Text ===\n")
                return text
                
        elif file_ext == '.docx':
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            print("\n=== Extracted Text from DOCX ===")
            print(text[:500] + "..." if len(text) > 500 else text)
            print("=== End of Extracted Text ===\n")
            return text
            
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
            
    except Exception as e:
        print(f"Error extracting text from file: {str(e)}")
        return ""

def prepare_vector_store(text: str):
    """Prepare vector store from text content."""
    try:
        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        chunks = text_splitter.split_text(text)
        
        # Create documents
        documents = [LangchainDocument(page_content=chunk) for chunk in chunks]
        
        # Create embeddings
        embeddings = OllamaEmbeddings(model="llama2")
        
        # Create vector store
        vector_store = FAISS.from_documents(documents, embeddings)
        
        return vector_store
    except Exception as e:
        print(f"Error preparing vector store: {str(e)}")
        return None

def generate_summary(vector_store) -> str:
    """Generate a concise summary of the uploaded content."""
    try:
        if vector_store is None:
            return "Please upload a document first."
            
        llm = ChatOllama(model="llama2", temperature=0.7)
        
        # Get all content from vector store
        all_docs = vector_store.similarity_search("", k=10)  # Get top chunks
        content = "\n".join([doc.page_content for doc in all_docs])
        
        prompt = f"""
        Please provide a concise summary of the following content. The summary should:
        1. Capture the main ideas and key points
        2. Be well-structured and easy to understand
        3. Be approximately 3-4 paragraphs long
        
        Content:
        {content}
        """
        
        response = llm.invoke(prompt)
        return response.content.strip()
        
    except Exception as e:
        print(f"Error generating summary: {str(e)}")
        return f"Error: {str(e)}"

def generate_study_notes(vector_store) -> str:
    """Generate structured study notes from the content."""
    try:
        if vector_store is None:
            return "Please upload a document first."
            
        llm = ChatOllama(model="llama2", temperature=0.7)
        
        # Get all content from vector store
        all_docs = vector_store.similarity_search("", k=10)  # Get top chunks
        content = "\n".join([doc.page_content for doc in all_docs])
        
        prompt = f"""
        Create clear and structured study notes from the following content. The notes should:
        1. Be organized with clear headings and subheadings
        2. Include bullet points for key concepts
        3. Highlight important terms or definitions
        4. Include examples where relevant
        5. Be easy to read and study from
        
        Content:
        {content}
        """
        
        response = llm.invoke(prompt)
        return response.content.strip()
        
    except Exception as e:
        print(f"Error generating study notes: {str(e)}")
        return f"Error: {str(e)}"

def get_ai_response(question: str, vector_store=None) -> str:
    """Get AI response based on the question and context."""
    try:
        if vector_store is None:
            return "Please upload a document first."
            
        # Initialize the LLM
        llm = ChatOllama(model="llama2", temperature=0.7)
        
        # Create QA chain
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vector_store.as_retriever(search_kwargs={"k": 3}),
            return_source_documents=True
        )
        
        # Get relevant context and prepare prompt
        context = vector_store.similarity_search(question, k=3)
        context_text = "\n".join([doc.page_content for doc in context])
        
        print("\n=== Retrieved Context ===")
        print(f"Question: {question}")
        print("\nContext:")
        print(context_text[:500] + "..." if len(context_text) > 500 else context_text)
        print("=== End of Context ===\n")
        
        prompt = f"""
        You are a helpful teaching assistant. Based on the following context and question, provide a clear and informative response.
        
        Context from course material:
        {context_text}
        
        Student Question:
        {question}
        
        Please provide a helpful response that:
        1. Directly addresses the question
        2. Uses information from the context
        3. Explains concepts clearly
        4. Provides examples if relevant
        """
        
        # Get response
        response = qa_chain({"query": prompt})
        
        if not response or 'result' not in response:
            return "Sorry, I couldn't generate a response. Please try again."
            
        return response['result'].strip()
            
    except Exception as e:
        print(f"Error getting AI response: {str(e)}")
        return f"Error: {str(e)}"

def store_query_response(question: str, response: str):
    """Store the query and response for future reference."""
    try:
        # Create a directory to store the query history if it doesn't exist
        history_dir = os.path.join(os.path.dirname(__file__), "query_history")
        os.makedirs(history_dir, exist_ok=True)
        
        # Create a timestamp-based filename
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = os.path.join(history_dir, f"query_{timestamp}.json")
        
        # Store the query and response
        data = {
            "timestamp": timestamp,
            "question": question,
            "response": response
        }
        
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
            
    except Exception as e:
        print(f"Failed to store query: {str(e)}")

def generate_mcq(topic: str, question: str, context: str) -> dict:
    """Generate a multiple choice question."""
    try:
        llm = ChatOllama(model="llama2", temperature=0.7)
        
        prompt = f"""
        Based on the following question and response about '{topic}', generate a multiple choice question.
        Make sure the question tests understanding of a key concept from the discussion.
        
        Question asked: {question}
        Response given: {context}
        
        Create a challenging but fair multiple choice question that:
        1. Tests understanding rather than mere memorization
        2. Has clear, unambiguous options
        3. Includes a thorough explanation for the correct answer
        
        Format your response EXACTLY like this, with no other text:
        {{
            "question": "The question text",
            "options": ["A) option1", "B) option2", "C) option3", "D) option4"],
            "correct_answer": "A",
            "explanation": "Brief explanation of why this is the correct answer"
        }}
        """
        
        response = llm.invoke(prompt).content
        
        # Clean the response
        response = response.strip()
        response = response.replace('```json', '').replace('```', '')
        response = response.strip()
        
        try:
            mcq = json.loads(response)
        except json.JSONDecodeError:
            # Try to fix common JSON formatting issues
            response = response.replace('\n', ' ').replace('\r', '')
            response = response.replace('"{', '{').replace('}"', '}')
            mcq = json.loads(response)
        
        # Validate MCQ format
        if not all(key in mcq for key in ["question", "options", "correct_answer", "explanation"]):
            raise ValueError("Missing required fields in MCQ")
        if len(mcq["options"]) != 4:
            raise ValueError("MCQ must have exactly 4 options")
        if mcq["correct_answer"] not in ["A", "B", "C", "D"]:
            raise ValueError("Invalid correct_answer format")
            
        return mcq
    except json.JSONDecodeError as e:
        print(f"JSON parsing error: {str(e)}")
        return {"error": "Failed to parse LLM response into valid MCQ format"}
    except Exception as e:
        print(f"Error generating MCQ: {str(e)}")
        return {"error": f"Error generating MCQ: {str(e)}"}

def generate_flashcards(topic: str, question: str, context: str) -> list:
    """Generate flashcards from the context."""
    try:
        llm = ChatOllama(model="llama2", temperature=0.7)
        
        prompt = f"""
        Based on the following question and response about '{topic}', generate 3 flashcards.
        Each flashcard should focus on a different key concept from the discussion.
        
        Question: {question}
        Response: {context}
        
        Create flashcards that:
        1. Are concise and clear
        2. Focus on important concepts
        3. Have clear, unambiguous answers
        4. Progress from basic to more complex concepts
        
        Format your response EXACTLY like this, with no other text:
        [
            {{"front": "Question 1", "back": "Answer 1"}},
            {{"front": "Question 2", "back": "Answer 2"}},
            {{"front": "Question 3", "back": "Answer 3"}}
        ]
        """
        
        response = llm.invoke(prompt).content
        
        # Clean the response
        response = response.strip()
        response = response.replace('```json', '').replace('```', '')
        response = response.strip()
        
        try:
            flashcards = json.loads(response)
        except json.JSONDecodeError:
            # Try to fix common JSON formatting issues
            response = response.replace('\n', ' ').replace('\r', '')
            response = response.replace('"{', '{').replace('}"', '}')
            flashcards = json.loads(response)
        
        # Validate flashcards
        if not isinstance(flashcards, list):
            raise ValueError("Response must be a list of flashcards")
        if len(flashcards) != 3:
            raise ValueError("Must generate exactly 3 flashcards")
        
        for card in flashcards:
            if not isinstance(card, dict) or 'front' not in card or 'back' not in card:
                raise ValueError("Invalid flashcard format")
            if not card['front'].strip() or not card['back'].strip():
                raise ValueError("Flashcard front and back cannot be empty")
        
        return flashcards
    except json.JSONDecodeError as e:
        print(f"JSON parsing error: {str(e)}")
        return {"error": "Failed to parse LLM response into valid flashcard format"}
    except Exception as e:
        print(f"Error generating flashcards: {str(e)}")
        return {"error": f"Error generating flashcards: {str(e)}"}

def load_instructor_settings():
    """Load the most recent instructor settings."""
    prompts_dir = "../instructor_prompts"
    if not os.path.exists(prompts_dir):
        return None
    
    # Get the most recent prompt file
    prompt_files = [f for f in os.listdir(prompts_dir) if f.startswith("instructor_settings_prompt_")]
    if not prompt_files:
        return None
    
    latest_prompt = max(prompt_files)
    with open(os.path.join(prompts_dir, latest_prompt), 'r') as f:
        return f.read() 