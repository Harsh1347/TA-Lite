import os
import fitz  # PyMuPDF
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from docx import Document as DocxDocument
import re

PDF_FOLDER = "teacher_data/materials"
VECTOR_DB_PATH = "teacher_data/faiss_index"
TRANSCRIPT_FOLDER = "teacher_data/transcript"

def clean_transcript_text(raw_text):
    cleaned_lines = []
    for line in raw_text.split("\n"):
        # Remove timestamps and numbers
        if re.match(r"^\d+$", line):
            continue
        if re.match(r"^\d\d:\d\d:\d\d\.\d+ -->", line):
            continue
        # Remove speaker names (e.g., "Jisun An:")
        line = re.sub(r"^\s*\w[\w\s]*:\s*", "", line)
        if line.strip():
            cleaned_lines.append(line.strip())
    return " ".join(cleaned_lines)


def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def load_and_split_docx(file_path):
    raw_text = extract_text_from_docx(file_path)
    cleaned_text = clean_transcript_text(raw_text)
    splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=150)
    chunks = splitter.split_text(cleaned_text)

    docs = []
    for i, chunk in enumerate(chunks):
        docs.append(Document(
            page_content=chunk,
            metadata={
                "source": os.path.basename(file_path),
                "chunk_id": i,
                "type": "transcript"
            }
        ))
    return docs


def extract_text_from_pdf(file_path):
    text = ""
    doc = fitz.open(file_path)
    for page in doc:
        text += page.get_text()
    return text

def load_and_split_pdf(file_path):
    raw_text = extract_text_from_pdf(file_path)
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    chunks = splitter.split_text(raw_text)

    docs = []
    for i, chunk in enumerate(chunks):
        docs.append(Document(
            page_content=chunk,
            metadata={
                "source": os.path.basename(file_path),
                "chunk_id": i,
                "type":"pdf"
            }
        ))
    return docs

def vectorize_all_documents():
    all_docs = []

    # Process PDFs
    for filename in os.listdir(PDF_FOLDER):
        if filename.endswith(".pdf"):
            file_path = os.path.join(PDF_FOLDER, filename)
            docs = load_and_split_pdf(file_path)
            all_docs.extend(docs)
            print(f"📄 Processed PDF: {filename} → {len(docs)} chunks")

    # Process DOCX transcripts
    for filename in os.listdir(TRANSCRIPT_FOLDER):
        if filename.endswith(".docx"):
            file_path = os.path.join(TRANSCRIPT_FOLDER, filename)
            docs = load_and_split_docx(file_path)
            all_docs.extend(docs)
            print(f"📝 Processed Transcript: {filename} → {len(docs)} chunks")

    if not all_docs:
        print("⚠️ No documents found to process.")
        return

    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    db = FAISS.from_documents(all_docs, embeddings)
    db.save_local(VECTOR_DB_PATH)
    print(f"\n✅ Vector store saved at: {VECTOR_DB_PATH}")

if __name__ == "__main__":
    vectorize_all_documents()